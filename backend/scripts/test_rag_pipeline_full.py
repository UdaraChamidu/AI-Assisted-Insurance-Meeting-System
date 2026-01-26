
import asyncio
import sys
import os
import io
import docx
from unittest.mock import MagicMock, AsyncMock, patch

# Add backend to path
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

# Add backend to path
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

async def test_full_rag_pipeline():
    print("🚀 Starting FULL Mock RAG Pipeline Verification...")
    
    # Patch config.settings GLOBALLY for this process
    with patch('config.settings') as mock_settings:
        mock_settings.app_env = 'test'
        mock_settings.ms_client_id = 'dummy'
        mock_settings.ms_client_secret = 'dummy'
        mock_settings.ms_tenant_id = 'dummy_tenant'
        mock_settings.pinecone_api_key = 'dummy'
        
        # Patch MSAL to prevent constructor error during import if any
        with patch('msal.ConfidentialClientApplication'):
            
            # Helper to create Services with mocked dependencies
            from services.ingestion_service import IngestionService
            
            # Create Mock Dependencies
            mock_sp = MagicMock()
            mock_emb = MagicMock()
            mock_pc = MagicMock()
            
            # Instantiate Service
            service = IngestionService()
            # INJECT MOCKS DIRECTLY
            service.sharepoint = mock_sp
            service.embeddings = mock_emb
            service.pinecone = mock_pc
            
            # 2. Simulate SharePoint Structure
            print("\n[1/5] Simulating SharePoint 'KB-DEV' Structure...")
    mock_sp.search_site.return_value = "site_kb_dev"
    mock_sp.get_drives.return_value = [{'name': 'Documents', 'id': 'drive_docs'}]
    
    # Mock folder mapping (Root -> Target Folders)
    mock_sp.list_drive_items.return_value = [
        {'name': '01_FL_State_Authority', 'id': 'folder_fl', 'folder': {}},
        {'name': '02_CMS_Medicare_Authority', 'id': 'folder_cms', 'folder': {}},
    ]
    
    # Mock recursive list for '02_CMS_Medicare_Authority' to find one PDF and one DOCX
    mock_sp.recursive_list_items.side_effect = lambda drive_id, folder_id=None: (
        [
            {
                'name': '2025_Medicare_Marketing_Guidelines.docx',
                'id': 'file_docx',
                'file': {},
                'fields': {'ProductUniverse': 'Medicare', 'AuthorityLevel': 'Federal'}
            }
        ] if folder_id == 'folder_cms' else []
    )
    
    # 3. Simulate File Content (DOCX)
    print("[2/5] Simulating DOCX Content Generation...")
    
    def create_mock_docx():
        file_stream = io.BytesIO()
        doc = docx.Document()
        doc.add_heading('Section 1: Permitted Marketing', level=1)
        doc.add_paragraph("Agents may only market compliant plans. " * 20) # ~100 words
        doc.add_heading('Section 2: Prohibited Activities', level=1)
        doc.add_paragraph("Agents must not use scare tactics. " * 30) # ~150 words
        doc.save(file_stream)
        return file_stream.getvalue()

    mock_sp.get_file_content.return_value = create_mock_docx()
    
    # 4. Run Pipeline
    print("[3/5] Running Ingestion Pipeline...")
    
    # Mock Embedding returning dummy vector
    mock_emb.embed_batch.side_effect = lambda texts: [[0.1] * 768] * len(texts)
    
    # Run
    await service.run_pipeline(site_name="KB-DEV")
    
    # 5. Verification
    print("\n[4/5] Verifying Metadata & Chunking...")
    
    # Verify Upsert Call
    if mock_pc.upsert_chunks.called:
        call_args = mock_pc.upsert_chunks.call_args[0][0] # First arg is list of vectors
        print(f"✅ Upsert called with {len(call_args)} chunks.")
        
        first_chunk = call_args[0]
        metadata = first_chunk['metadata']
        
        print(f"   Chunk ID: {first_chunk['id']}")
        print(f"   Metadata Universe: {metadata.get('universe')}")
        print(f"   Metadata Product: {metadata.get('product_universe')}")
        
        # ASSERTIONS
        assert metadata.get('universe') == '02_CMS_Medicare_Authority', "❌ Failed: Universe tag missing or incorrect!"
        assert metadata.get('product_universe') == 'Medicare', "❌ Failed: ProductUniverse metadata missing!"
        assert "Meeting" not in metadata, "❌ Failed: Unexpected metadata found."
        
        print("✅ Metadata validation passed!")
    else:
        print("❌ Upsert NOT called!")

    print("\n[5/5] RAG Pipeline Verification Complete!")

if __name__ == "__main__":
    asyncio.run(test_full_rag_pipeline())
